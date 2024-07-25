import os
import fire
from docx import Document

# DEMO = (
#     'You can create a new word file by calling `word_read_file` with 1 arguments.\n'
#     '1. The path of the word file: file_path: str\n'
#     "You can call it by: {'app': 'word', 'action': 'word_read_file', 'file_path': ...}"
# )

DEMO = (
    "read the content of the word file: "
    "{'app': 'word', 'action': 'read_file', 'file_path': [THE_PATH_TO_THE_WORD_FILE]"
)


def construct_action(work_dir, args: dict, py_file_path='/apps/word_app/word_read_file.py'):
    # TODO: not sure if we need to specify the file path with the current workdir
    return f'python3 {py_file_path} --file_path {args["file_path"]}'

def word_read_file(file_path):
    doc = Document(file_path)
    return doc.paragraphs

def word_read_file_into_string(file_path):
    doc = Document(file_path)
    string = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return string
        
def wrap(paragraphs):
    ob = ("The following is the content from the word file:")
    for paragraph in paragraphs:
        ob += f"\n{paragraph.text}"
    return ob

def main(file_path, debug=False):
    if not os.path.exists(file_path):
        return f"OBSERVATION: The file {file_path} does not exist. Failed to read the file."
    paragraphs = word_read_file(file_path)
    if debug:
        print(paragraphs)
    observation = wrap(paragraphs)
    return 'OBSERVATION: ' + observation


if __name__ == '__main__':
    fire.Fire(main)