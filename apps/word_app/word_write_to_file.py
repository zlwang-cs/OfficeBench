import os
import fire
import pandas as pd
from docx import Document

# DEMO = (
#     'You can create a new word file by calling `word_write_to_file` with 2 arguments.\n'
#     '1. The path of the word file: file_path: str\n'
#     '2. The content you wish to write to the word file: contents: str\n'
#     "You can call it by: {'app': 'word', 'action': 'word_read_file', 'file_path': ..., 'contents': ...}"
# )

DEMO = (
    "write text to a word file: "
    "{'app': 'word', 'action': 'write_to_file', 'file_path': [THE_PATH_TO_THE_WORD_FILE], 'contents': [THE_CONTENTS_YOU_WISH_TO_WRITE]}"
)




def construct_action(work_dir, args: dict, py_file_path='/apps/word_app/word_write_to_file.py'):
    # TODO: not sure if we need to specify the file path with the current workdir
    # return f'python3 {py_file_path} --file_path {args["file_path"]} --contents "{args["contents"]}"'
    comments = args['contents'].replace("'", "").replace('"', '')
    return "python3 {} --file_path {} --contents '''{}''' ".format(py_file_path, args['file_path'], comments)






def word_write_to_file(file_path, contents: str, style='pure-text'):
    try:
        if os.path.exists(file_path):
            document = Document(file_path)
        else:
            document = Document()
        if style == 'pure-text':
            document.add_paragraph(contents)
        elif style == 'title':
            document.add_heading(contents, 0)
        elif style == 'subtitle':
            document.add_heading(contents, 1)
        else:
            raise NotImplementedError
        document.save(file_path)
        return True
    except:
        return False


def main(file_path, contents, style='pure-text'):
    if not os.path.exists(file_path):
        return f'OBSERVATION: The file {file_path} does not exist. Failed to write to the file.'
    
    success = word_write_to_file(file_path, contents, style)
    if success:
        observation = f'Successfully write contents to {file_path}'
    else:
        observation = f'Failed to write contents to {file_path}'
    return 'OBSERVATION: ' + observation

if __name__ == '__main__':
    fire.Fire(main)
